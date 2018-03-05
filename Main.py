import os
import sys

import json
import cv2
import numpy as np
import easygui
from PyQt5 import QtWidgets, QtCore, QtGui
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtGui import QIntValidator
from PyQt5.uic import loadUi
from scipy.ndimage import imread
from docx import *
from datetime import datetime

os.environ["QT_QPA_PLATFORM_PLUGIN_PATH"] = "C:\Python352\Lib\site-packages\PyQt5\plugins\platforms"
if not os.path.exists("Storage"):
    os.makedirs('Storage')


class SaveWin(QtWidgets.QMainWindow):
    def __init__(self, output_text, contours, processed_image):
        QtWidgets.QMainWindow.__init__(self)
        super(SaveWin, self).__init__()
        loadUi('SaveInfo.ui', self)
        self.setWindowTitle('Save')
        self.pushButton.clicked.connect(self.save)
        self.pushButton_2.clicked.connect(self.exit)
        self.processed_image = processed_image
        self.output_text = output_text
        self.contours = contours
        self.full_name = self.lineEdit
        self.birthday = self.dateEdit_2
        self.attendance = self.dateEdit
        self.description = self.textEdit

    def save(self):
        print(self.full_name.text())
        print(self.processed_image)
        os.makedirs('Storage/'+str(self.full_name.text()))
        doc = Document()
        temp = cv2.imread("C:/Users/Marko/Documents/ImageStorage/temp.jpg")
        cv2.imwrite('Storage/'+str(self.full_name.text())+'/'+str(self.attendance.date().toPyDate()) + '.jpg', temp)
        doc.add_heading(str(self.full_name.text())+" " + str(self.attendance.date().toPyDate())).bold = True
        doc.add_paragraph("Birthday:"+str(self.birthday.date().toPyDate()))
        doc.add_heading("Charasteristics").bold = True
        doc.add_paragraph(self.output_text)
        doc.add_heading("Description:").bold = True
        doc.add_paragraph(str(self.description.toPlainText()))
        doc.save('Storage/'+str(self.full_name.text())+'/'+str(self.attendance.date().toPyDate()) + '.docx')
        self.close()

    def exit(self):
        self.close()


class ProcessedWin(QtWidgets.QMainWindow):
    def __init__(self, path, x, y):
        QtWidgets.QMainWindow.__init__(self)
        super(ProcessedWin, self).__init__()
        loadUi('Result.ui', self)
        self.pushButton.clicked.connect(self.goon)
        self.pushButton_2.clicked.connect(self.save)
        self.output_text = self.textEdit
        self.dimensional_area = x*y
        self.contours = []
        imagelabel = self.labelBefore
        imagelabel_after = self.labelAfter
        input_image = imread(path)
        height, width, channels = input_image.shape
        bytesPerLine = channels * width
        qImg = QtGui.QImage(input_image.data, width, height, bytesPerLine, QtGui.QImage.Format_RGB888)
        pixmap01 = QtGui.QPixmap.fromImage(qImg)
        pixmap_image = QtGui.QPixmap(pixmap01)
        imagelabel.setPixmap(pixmap_image)
        imagelabel.setAlignment(QtCore.Qt.AlignCenter)
        imagelabel.setScaledContents(True)
        imagelabel.setMinimumSize(1, 1)
        imagelabel.show()

        processing_image = cv2.imread(path)
        response = processing_image
        response = cv2.cvtColor(response, cv2.COLOR_BGR2HSV)
        font = cv2.FONT_HERSHEY_SIMPLEX
        # loop over the boundaries
        # create NumPy arrays from the boundaries
        lower_blue = np.array([110, 50, 50], dtype=np.uint8)
        upper_blue = np.array([130, 255, 255], dtype=np.uint8)

        # find the colors within the specified boundaries and apply
        # the mask
        mask = cv2.inRange(response, lower_blue, upper_blue)
        output = cv2.bitwise_and(processing_image, processing_image, mask=mask)
        gray = cv2.cvtColor(output, cv2.COLOR_BGR2GRAY);
        gray = cv2.GaussianBlur(gray, (9, 9), 0)
        _, bin = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)  # inverted threshold (light obj on dark bg)
        bin = cv2.dilate(bin, None)  # fill some holes
        bin = cv2.dilate(bin, None)
        bin = cv2.erode(bin, None)  # dilate made our shape larger, revert that
        bin = cv2.erode(bin, None)
        bin, contours1, hierarchy1 = cv2.findContours(bin, cv2.RETR_CCOMP, cv2.CHAIN_APPROX_SIMPLE)
        hie = hierarchy1[0]
        self.output_text_val = ""
        z = 1
        for x in zip(contours1, hie):
            currentContour = x[0]
            currentHierarchy = x[1]
            if currentHierarchy[2] < 0:
                epsilon = 0.1 * cv2.arcLength(currentContour, True)
                approx = cv2.approxPolyDP(currentContour, epsilon, True)
                hull = cv2.convexHull(currentContour)
                current = Object()
                current.contour = currentContour
                current.area = (self.dimensional_area*cv2.contourArea(currentContour)/bin.size)
                current.equi_diameter = np.sqrt(4*current.area/np.pi)
                (x, y), (MA, ma), angle = cv2.fitEllipse(current.contour)
                current.major_axis = MA
                current.minor_axis = ma
                Momentum = cv2.moments(currentContour)
                cx = int(Momentum['m10'] / Momentum['m00'])
                cy = int(Momentum['m01'] / Momentum['m00'])
                current.momentum = Momentum
                current.center = (cx, cy)
                current.contour_number = z
                current.perimeter = cv2.arcLength(current.contour, True)
                cv2.circle(processing_image, (cx, cy), 3, (255, 0, 255), -1)
                cv2.putText(processing_image, str(z) + 'C', (cx, cy), font, 0.5, (255, 0, 255), 2, cv2.LINE_AA)
                current.leftmost = tuple(currentContour[currentContour[:, :, 0].argmin()][0])
                current.rightmost = tuple(currentContour[currentContour[:, :, 0].argmax()][0])
                current.topmost = tuple(currentContour[currentContour[:, :, 1].argmin()][0])
                current.bottommost = tuple(currentContour[currentContour[:, :, 1].argmax()][0])
                cv2.circle(processing_image, current.leftmost, 5, (255, 0, 255), -1)
                cv2.circle(processing_image, current.rightmost, 5, (255, 0, 255), -1)
                cv2.circle(processing_image, current.topmost, 5, (255, 0, 255), -1)
                cv2.circle(processing_image, current.bottommost, 5, (255, 0, 255), -1)
                processing_image = cv2.drawContours(processing_image, currentContour, -1, (0, 0, 255), 2)
                self.contours.append(current)
                z = z + 1
                self.output_text_val += "Contour" + str(current.contour_number) + ":" + '\n' + "Perimenter:" + str(current.perimeter) + '\n' + "Area(cm^2):" + str(current.area) + '\n'
                self.output_text_val += "Equivalent diameter:" + str(current.equi_diameter) + '\n'

        self.output_text.setText(self.output_text_val)
        imageforshow = processing_image
        height, width, channels = imageforshow.shape
        bytesPerLine = channels * width
        qImg = QtGui.QImage(imageforshow.data, width, height, bytesPerLine, QtGui.QImage.Format_RGB888)
        pixmap01 = QtGui.QPixmap.fromImage(qImg)
        pixmap_image = QtGui.QPixmap(pixmap01)
        imagelabel_after.setPixmap(pixmap_image)
        imagelabel_after.setAlignment(QtCore.Qt.AlignCenter)
        imagelabel_after.setScaledContents(True)
        imagelabel_after.setMinimumSize(1, 1)
        imagelabel_after.show()
        dt = datetime.today()
        cv2.imwrite("C:/Users/Marko/Documents/ImageStorage/temp.jpg", processing_image)
        self.processed_image = "C:/Users/Marko/Documents/ImageStorage/temp.jpg"

    def goon(self):
        self.close()

    def save(self):
        self.dialog = SaveWin(self.output_text.toPlainText(), self.contours, self.processed_image)
        self.dialog.show()


class MyWin(QtWidgets.QMainWindow):
    def __init__(self):
        QtWidgets.QMainWindow.__init__(self)
        super(MyWin, self).__init__()
        loadUi('Base.ui', self)
        self.setWindowTitle('TestDrive')
        self.pushButtonUpload.clicked.connect(self.upload_click)
        self.pushButton_2.clicked.connect(self.process_click)
        self.pushButton.clicked.connect(self.exit)
        self.digit_input_x = self.lineEdit
        self.digit_input_y = self.lineEdit_2
        self.onlyInt = QIntValidator()
        self.digit_input_x.setValidator(self.onlyInt)
        self.digit_input_y.setValidator(self.onlyInt)

    @pyqtSlot()
    def upload_click(self):
        imagelabel = self.label_5
        filepath = easygui.fileopenbox()
        print(filepath)
        if not filepath:
            return 0
        input_image = imread(filepath)
        height, width, channels = input_image.shape
        bytesPerLine = channels * width
        qImg = QtGui.QImage(input_image.data, width, height, bytesPerLine, QtGui.QImage.Format_RGB888)
        pixmap01 = QtGui.QPixmap.fromImage(qImg)
        pixmap_image = QtGui.QPixmap(pixmap01)
        imagelabel.setPixmap(pixmap_image)
        imagelabel.setAlignment(QtCore.Qt.AlignCenter)
        imagelabel.setScaledContents(True)
        imagelabel.setMinimumSize(1, 1)
        imagelabel.show()
        self.pathforsend = filepath
    def process_click(self):
        pth = self.imagepath()
        if not pth:
            return 0
        else:
            x = self.digit_input_x.text()
            y = self.digit_input_y.text()
            x = int(x)
            y = int(y)
            self.dialog = ProcessedWin(pth, x, y)
            self.dialog.show()

    def imagepath(self):
        return self.pathforsend

    def exit(self):
        self.close()


class Object:
    def toJSON(self):
        return json.dumps(self, default=lambda o: o.__dict__,sort_keys=True, indent=4)


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    widget = MyWin()
    widget.show()
    sys.exit(app.exec_())
    widget.show(sys.exit(app.exec_()))
